#!/usr/bin/env python
# _*_ coding:utf-8 _*_
from docx import Document
from docx.shared import Inches
import os

# ''' 创建输出路径  '''
# cwd = os.getcwd()
# print(cwd)
# global outputFolderName
# outputFolderName = cwd + r'\out'  # + datetime.datetime.now().strftime('%Y.%m.%d.%H.%M')
# outPathExists = os.path.exists(outputFolderName)
# if not outPathExists:
#     os.makedirs(outputFolderName)
#     print(outputFolderName + ' 路径创建成功')
# else:
#     print(outputFolderName + ' 路径已存在')


document = Document()
document.add_heading('计算报告', 0)

p = document.add_paragraph('第一行 ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

document.add_heading('输入参数', level=1)
# document.add_paragraph(
#      '极对数', style='List Number'
#  )

table_para = document.add_table(rows=6, cols=2)
#print(dir(table_para))
hdr_cells = table_para.columns[0].cells

hdr_cells[0].text = '极对数'
hdr_cells[1].text = '上网总功率[kW]'
hdr_cells[2].text = '转子开口电压[Vrms]'
hdr_cells[3].text = '电网电压[Vrms]'
hdr_cells[4].text = '电网频率[Hz]'
hdr_cells[5].text = '定子无功功率[kW]'




records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)


document.add_heading('输入功率曲线表', level=1)
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '转速[RPM]'
hdr_cells[1].text = '上网功率[kW]'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

# document.add_paragraph('Intense quote', style='Intense Quote')
#
# document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='List Number'
# )
#
# document.add_picture('monty-truth.png', width=Inches(1.25))



document.add_page_break()
document.save('demo.docx')

'''

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

document.add_picture('monty-truth.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('demo.docx')
'''